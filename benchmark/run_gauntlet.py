from __future__ import annotations

import io
import json
import math
import os
import statistics
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal

import httpx
from docx import Document
from PIL import Image, ImageDraw

BASE_URL = os.getenv("MILITARY_SLICES_BENCHMARK_URL", "https://competition-rc---military-slices-ztvqlzospa-uw.a.run.app")
REVISION = os.getenv("MILITARY_SLICES_BENCHMARK_REVISION", "military-slices-00016-miz")
SOURCE_COMMIT = os.getenv("MILITARY_SLICES_BENCHMARK_COMMIT", "b200d5b940b058d1cd6e805c88acd47ed098835b")
STARTING_ESTIMATED_SPEND = 0.0
# Deliberately conservative benchmark accounting rates; actual provider billing was not exposed.
INPUT_USD_PER_MILLION = 2.0
OUTPUT_USD_PER_MILLION = 8.0
TOOL_CALL_USD = 0.01
SAFE_COST_RAIL = 27.50
HARD_COST_RAIL = 30.00

ActionKind = Literal["text", "artifact", "reject", "revalidate", "reload", "replay", "decide"]


@dataclass(frozen=True)
class Action:
    kind: ActionKind
    value: str = ""
    artifact_type: str = ""


@dataclass(frozen=True)
class Scenario:
    id: str
    service: str
    persona: str
    stage: str
    family: str
    actions: tuple[Action, ...]
    expect_anchor: bool = True
    forbidden_slices: tuple[str, ...] = ()
    expected: tuple[str, ...] = ()
    max_model_calls: int = 4


SCENARIOS: tuple[Scenario, ...] = (
    Scenario("S01", "Army", "separating member", "12–18 months", "straight-line", (Action("text", "I leave the Army in December 2027. My career target is Logistics Program Manager. I will stay local near Fort Cavazos."),), expected=("target", "task_bound")),
    Scenario("S02", "Navy", "retiring member", "6–12 months", "messy-first-input", (Action("text", "I retire from the Navy in June 2027. I need steady civilian work, my spouse teaches locally, I cannot decide between operations and project work, and I also coach youth baseball."),), expected=("anchor", "task_bound")),
    Scenario("S03", "Marine Corps", "separating member", ">18 months", "no-goal", (Action("text", "I am a Marine with a family and a lot to think about after service, but I do not know what I want yet."),), expect_anchor=False, expected=("anchor_gate",)),
    Scenario("S04", "Air Force", "separating member", "90–180 days", "closest-equivalent", (Action("text", "I separate from the Air Force in January 2027 and want the closest civilian equivalent to aircraft maintenance production supervision. I can relocate."),), expected=("anchor", "task_bound")),
    Scenario("S05", "Space Force", "separating member", "12–18 months", "education-first", (Action("text", "I leave the Space Force in March 2028. Education comes first; I want a graduate certificate before testing space systems program management."),), expected=("anchor", "task_bound")),
    Scenario("S06", "Coast Guard", "recent veteran", "recently separated", "post-transition-reevaluation", (Action("text", "I recently left the Coast Guard and my career target is Emergency Management Coordinator. I need to stay near Seattle for family."), Action("reload")), expected=("target", "persistence")),
    Scenario("S07", "Army", "military spouse", "PCS-driven", "spouse-pcs-licensing", (Action("text", "My spouse is Army and we PCS to Colorado Springs in eight months. My anchor is protecting my nursing career and resolving Colorado licensing before the move."),), expected=("anchor", "task_bound")),
    Scenario("S08", "Air Force", "dual-military household", "6–12 months", "family-geography", (Action("text", "We are dual Air Force and one of us separates in September 2027. Childcare and school stability mean we will stay local; employment comes first."),), expected=("anchor", "task_bound")),
    Scenario("S09", "Navy", "retiring member", "6–12 months", "resume-scope-drift", (Action("text", "My anchor is make my resume submission-ready for Program Manager roles."), Action("artifact", artifact_type="docx"), Action("reload")), forbidden_slices=("Education", "Location"), expected=("persistence", "task_bound")),
    Scenario("S10", "Marine Corps", "separating member", "<90 days", "pdf-evidence", (Action("text", "My career target is Operations Manager and I separate in October 2026."), Action("artifact", artifact_type="pdf")), expected=("target", "task_bound")),
    Scenario("S11", "Space Force", "military spouse", "PCS-driven", "image-evidence", (Action("text", "We PCS next spring and my target role is Technical Project Coordinator."), Action("artifact", artifact_type="png")), expected=("target", "task_bound")),
    Scenario("S12", "Coast Guard", "separating member", "90–180 days", "multiple-artifacts", (Action("text", "My anchor is get a federal emergency planning job before separation."), Action("artifact", artifact_type="txt"), Action("artifact", artifact_type="docx")), expected=("anchor", "task_bound")),
    Scenario("S13", "Navy", "retiring member", "6–12 months", "target-change-location-revalidation", (Action("text", "I retire in June 2027. My career target is Program Management. I will stay local."), Action("text", "My career target is Defense Aerospace Program Management."), Action("revalidate"), Action("reload")), forbidden_slices=("Education",), expected=("impact", "revalidation", "persistence")),
    Scenario("S14", "Army", "separating member", "12–18 months", "unaffected-education", (Action("text", "My career target is Supply Chain Manager. My minimum compensation is 90000 dollars and I will stay local."), Action("text", "My minimum compensation is now 100000 dollars.")), forbidden_slices=("Education",), expected=("no_education_impact",)),
    Scenario("S15", "Air Force", "separating member", "6–12 months", "separation-date-fanout", (Action("text", "I separate in June 2027. My application deadline is April 2027, education starts August 2027, and I plan to relocate in July 2027."), Action("text", "My separation date changed to December 2027.")), expected=("single_impact", "task_bound")),
    Scenario("S16", "Coast Guard", "separating member", "90–180 days", "stale-before-conflict", (Action("text", "My career target is Port Operations Manager. I will stay local."), Action("text", "My career target is National Maritime Operations Manager and that role may require relocation.")), expected=("impact", "no_false_conflict", "no_false_paralysis")),
    Scenario("S17", "Marine Corps", "separating member", "<90 days", "genuine-conflict", (Action("text", "I need immediate income and I also require full-time education before I will accept work."),), expected=("conflict",)),
    Scenario("S18", "Army", "separating member", "terminal period", "human-authority", (Action("text", "I am on terminal leave and I need help choosing whether employment or full-time school is my first priority."),), expected=("human_gate",)),
    Scenario("S19", "Navy", "retiring member", "6–12 months", "rejection", (Action("text", "I retire next year and want civilian operations work but I am not sure which role."), Action("reject"), Action("reload")), expected=("rejection", "persistence", "task_bound")),
    Scenario("S20", "Air Force", "recent veteran", "post-transition", "changed-answer-replay", (Action("text", "My career target is Project Coordinator. I can relocate."), Action("replay"), Action("text", "I will now stay local because of family."), Action("reload")), expected=("idempotency", "persistence")),
    Scenario("S21", "Space Force", "separating member", ">18 months", "entrepreneurship", (Action("text", "I leave the Space Force in 2028 and my anchor is build a small cybersecurity consulting business, not find a traditional job."),), expected=("anchor", "task_bound")),
    Scenario("S22", "Marine Corps", "veteran", "post-transition", "career-change", (Action("text", "I was a Marine logistics specialist, but my career target is High School History Teacher. Education first, and I will not relocate."),), expected=("target", "task_bound")),
)


def artifact_payload(kind: str) -> tuple[str, bytes, str]:
    text = "Synthetic evidence: led 18 people, managed schedules, controlled resources, and briefed leaders. Career target remains Program Manager."
    if kind == "txt":
        return "synthetic-resume.txt", text.encode(), "text/plain"
    if kind == "docx":
        doc = Document()
        doc.add_heading("Synthetic Transition Resume", 1)
        doc.add_paragraph(text)
        doc.add_paragraph("Additional cyber and maintenance evidence is background only unless the declared target needs it.")
        buf = io.BytesIO()
        doc.save(buf)
        return "synthetic-resume.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if kind == "pdf":
        image = Image.new("RGB", (1200, 1600), "white")
        draw = ImageDraw.Draw(image)
        draw.multiline_text((70, 70), text + "\nTarget evidence: operations planning and cross-team delivery.", fill="black", spacing=18)
        buf = io.BytesIO()
        image.save(buf, format="PDF", resolution=150)
        return "synthetic-resume.pdf", buf.getvalue(), "application/pdf"
    image = Image.new("RGB", (1000, 500), "white")
    draw = ImageDraw.Draw(image)
    draw.multiline_text((35, 35), text + "\nSynthetic LinkedIn screenshot.", fill="black", spacing=12)
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return "synthetic-linkedin.png", buf.getvalue(), "image/png"


def post_json(client: httpx.Client, path: str, body: dict[str, Any]) -> tuple[dict[str, Any], float]:
    started = time.perf_counter()
    response = client.post(path, json=body)
    latency = time.perf_counter() - started
    response.raise_for_status()
    return response.json(), latency


def metrics(state: dict[str, Any]) -> dict[str, int | float]:
    return state.get("telemetry", {})


def run_scenario(scenario: Scenario) -> dict[str, Any]:
    started = time.perf_counter()
    failures: list[str] = []
    events: list[dict[str, Any]] = []
    latencies: list[dict[str, Any]] = []
    all_states: list[dict[str, Any]] = []
    model_backed_latencies: list[float] = []
    deterministic_latencies: list[float] = []
    last_confirm_body: dict[str, Any] | None = None
    last_impact: dict[str, Any] | None = None
    max_tasks = 0
    max_impacts = 0
    with httpx.Client(base_url=BASE_URL, timeout=45, follow_redirects=True) as client:
        first_started = time.perf_counter()
        envelope = client.get("/api/state").raise_for_status().json()
        latencies.append({"kind": "state", "seconds": time.perf_counter() - first_started, "model": False})
        initial = envelope["state"]
        initial_version = initial["version"]
        receipt_start = len(json.dumps(initial, separators=(",", ":"), sort_keys=True).encode())
        all_states.append(initial)
        for index, action in enumerate(scenario.actions):
            before_version = envelope["state"]["version"]
            if action.kind == "text":
                orientation, orient_latency = post_json(client, "/api/orient", {"text": action.value})
                latencies.append({"kind": "orient", "seconds": orient_latency, "model": False})
                untouched = client.get("/api/state").raise_for_status().json()["state"]["version"]
                if untouched != before_version:
                    failures.append("UNAUTHORIZED_MUTATION")
                key = f"{scenario.id.lower()}-{index}-{uuid.uuid4().hex}"
                body = {
                    "token": orientation["token"],
                    "reviewed_input": orientation["reviewed_input"],
                    "expected_version": before_version,
                    "idempotency_key": key,
                }
                envelope, latency = post_json(client, "/api/confirm", body)
                used_model = bool(envelope.get("agent_run") and envelope["agent_run"].get("provider") != "deterministic")
                latencies.append({"kind": "confirm", "seconds": latency, "model": used_model})
                (model_backed_latencies if used_model else deterministic_latencies).append(latency)
                last_confirm_body = body
                events.append({"action": "text", "input": action.value, "orientation": orientation, "version": envelope["state"]["version"], "agent_run": envelope.get("agent_run")})
            elif action.kind == "artifact":
                filename, data, content_type = artifact_payload(action.artifact_type)
                request_started = time.perf_counter()
                response = client.post(
                    "/api/artifact",
                    data={"expected_version": str(before_version), "idempotency_key": f"{scenario.id.lower()}-artifact-{index}-{uuid.uuid4().hex}"},
                    files={"file": (filename, data, content_type)},
                )
                latency = time.perf_counter() - request_started
                response.raise_for_status()
                envelope = response.json()
                used_model = bool(envelope.get("agent_run") and envelope["agent_run"].get("provider") != "deterministic")
                latencies.append({"kind": f"artifact:{action.artifact_type}", "seconds": latency, "model": used_model})
                (model_backed_latencies if used_model else deterministic_latencies).append(latency)
                events.append({"action": "artifact", "type": action.artifact_type, "version": envelope["state"]["version"], "agent_run": envelope.get("agent_run")})
            elif action.kind == "reject":
                hypotheses = envelope["state"].get("career_hypotheses", [])
                active = envelope.get("active_gate")
                if not hypotheses or not active or active.get("id") != "career-direction":
                    failures.append("OTHER")
                    events.append({"action": "reject", "skipped": "no career hypothesis gate"})
                    continue
                title = hypotheses[0]["title"]
                envelope, latency = post_json(client, "/api/decision", {"gate_id": "career-direction", "value": f"reject:{title}", "expected_version": before_version, "idempotency_key": f"{scenario.id.lower()}-reject-{uuid.uuid4().hex}"})
                used_model = bool(envelope.get("agent_run") and envelope["agent_run"].get("provider") != "deterministic")
                latencies.append({"kind": "reject", "seconds": latency, "model": used_model})
                (model_backed_latencies if used_model else deterministic_latencies).append(latency)
                events.append({"action": "reject", "title": title, "version": envelope["state"]["version"]})
            elif action.kind == "revalidate":
                impact = envelope.get("impact")
                if not impact:
                    failures.append("BAD_REVALIDATION")
                    events.append({"action": "revalidate", "skipped": "no impact"})
                    continue
                last_impact = impact
                envelope, latency = post_json(client, "/api/revalidate", {"impact_id": impact["id"], "action": "confirm", "expected_version": before_version, "idempotency_key": f"{scenario.id.lower()}-revalidate-{uuid.uuid4().hex}"})
                deterministic_latencies.append(latency)
                latencies.append({"kind": "revalidate", "seconds": latency, "model": False})
                events.append({"action": "revalidate", "impact": impact, "version": envelope["state"]["version"]})
            elif action.kind == "decide":
                gate = envelope.get("active_gate")
                if not gate:
                    failures.append("OTHER")
                    events.append({"action": "decide", "skipped": "no active gate"})
                    continue
                envelope, latency = post_json(
                    client,
                    "/api/decision",
                    {
                        "gate_id": gate["id"],
                        "value": action.value,
                        "expected_version": before_version,
                        "idempotency_key": f"{scenario.id.lower()}-decide-{uuid.uuid4().hex}",
                    },
                )
                deterministic_latencies.append(latency)
                latencies.append({"kind": "decide", "seconds": latency, "model": False})
                events.append({"action": "decide", "gate_id": gate["id"], "version": envelope["state"]["version"]})
            elif action.kind == "reload":
                request_started = time.perf_counter()
                reloaded = client.get("/api/state").raise_for_status().json()
                latency = time.perf_counter() - request_started
                latencies.append({"kind": "reload", "seconds": latency, "model": False})
                if reloaded["state"]["version"] != before_version or reloaded["state"].get("human_anchor") != envelope["state"].get("human_anchor"):
                    failures.append("PERSISTENCE_FAILURE")
                envelope = reloaded
                events.append({"action": "reload", "version": envelope["state"]["version"]})
            elif action.kind == "replay":
                if not last_confirm_body:
                    failures.append("IDEMPOTENCY_FAILURE")
                    continue
                replay, latency = post_json(client, "/api/confirm", last_confirm_body)
                latencies.append({"kind": "replay", "seconds": latency, "model": False})
                if replay["state"]["version"] != before_version:
                    failures.append("IDEMPOTENCY_FAILURE")
                envelope = replay
                events.append({"action": "replay", "version": envelope["state"]["version"]})
            state = envelope["state"]
            all_states.append(state)
            max_tasks = max(max_tasks, len(state.get("active_tasks", [])))
            max_impacts = max(max_impacts, len(state.get("impacts", [])))
            if len(state.get("active_tasks", [])) > 3:
                failures.append("TASK_OVERFLOW")
            if metrics(state).get("temporal_full_rebuilds", 0):
                failures.append("FULL_REBUILD")
            if metrics(state).get("temporal_freshness_model_calls", 0):
                failures.append("MODEL_OVERUSE")

        final = envelope["state"]
        final_metrics = metrics(final)
        if scenario.expect_anchor and not final.get("human_anchor"):
            failures.append("ANCHOR_DRIFT")
        if not scenario.expect_anchor and final.get("human_anchor") and "do not know" not in (final.get("human_anchor") or "").lower():
            failures.append("UNAUTHORIZED_ACTIVATION")
        if max_tasks > 3:
            failures.append("TASK_OVERFLOW")
        if len(json.dumps(final, separators=(",", ":"), sort_keys=True).encode()) > 200_000:
            failures.append("RECEIPT_BLOAT")
        if final_metrics.get("model_calls", 0) > scenario.max_model_calls:
            failures.append("MODEL_OVERUSE")
        if "impact" in scenario.expected and max_impacts == 0:
            failures.append("BAD_REVALIDATION")
        if "revalidation" in scenario.expected and (last_impact is None or final_metrics.get("temporal_one_tap_confirmations", 0) < 1):
            failures.append("BAD_REVALIDATION")
        if "no_false_conflict" in scenario.expected and final.get("conflicts"):
            failures.append("FALSE_CONFLICT")
        execution_state = final.get("execution", {}).get("state")
        if "no_false_paralysis" in scenario.expected and execution_state == "PARALYZED":
            failures.append("FALSE_PARALYSIS")
        if "conflict" in scenario.expected and not final.get("conflicts"):
            failures.append("MISSED_CONFLICT")
        if "paralyzed" in scenario.expected and execution_state != "PARALYZED":
            failures.append("MISSED_PARALYSIS")
        if "complete" in scenario.expected and execution_state != "COMPLETE":
            failures.append("MISSED_COMPLETION")
        if "active" in scenario.expected and execution_state != "ACTIVE":
            failures.append("FALSE_PARALYSIS")
        if "resume_target_gate" in scenario.expected and not any(
            gate.get("id") == "resume-target-role" for gate in final.get("gates", [])
        ):
            failures.append("MISSED_RESUME_TARGET_GATE")
        if "no_resume_target_gate" in scenario.expected and any(
            gate.get("id") == "resume-target-role" for gate in final.get("gates", [])
        ):
            failures.append("FALSE_RESUME_TARGET_CLOSURE")
        if "scoped_evidence" in scenario.expected and not any(
            "led 20 people" in fact.get("statement", "").casefold() for fact in final.get("facts", [])
        ):
            failures.append("SCOPED_CONTINUATION_FAILURE")
        if "rejection" in scenario.expected and not final.get("rejected_roles"):
            failures.append("OTHER")
        if "single_impact" in scenario.expected and max_impacts > 1:
            failures.append("BAD_REVALIDATION")
        if "no_education_impact" in scenario.expected and any(item.get("affected_slice") == "Education" for item in final.get("impacts", [])):
            failures.append("IRRELEVANT_SLICE_ACTIVATION")
        receipt_end = len(json.dumps(final, separators=(",", ":"), sort_keys=True).encode())
        input_tokens = int(final_metrics.get("input_tokens", 0))
        output_tokens = int(final_metrics.get("output_tokens", 0))
        tool_calls = int(final_metrics.get("tool_calls", 0))
        estimated_cost = input_tokens / 1_000_000 * INPUT_USD_PER_MILLION + output_tokens / 1_000_000 * OUTPUT_USD_PER_MILLION + tool_calls * TOOL_CALL_USD
        return {
            "scenario_id": scenario.id,
            "persona_type": scenario.persona,
            "service": scenario.service,
            "transition_stage": scenario.stage,
            "scenario_family": scenario.family,
            "human_anchor": final.get("human_anchor"),
            "starting_state_version": initial_version,
            "turns": len(scenario.actions),
            "artifacts_used": [a.artifact_type for a in scenario.actions if a.kind == "artifact"],
            "gates_evaluated": len({gate["id"] for state in all_states for gate in state.get("gates", [])}),
            "gates_closed_machine": int(final_metrics.get("agent_gates_closed", 0)),
            "human_gates": 1 if envelope.get("active_gate") else 0,
            "conflicts": final.get("conflicts", []),
            "paralyzed_transitions": 1 if execution_state == "PARALYZED" else 0,
            "active_tasks_max": max_tasks,
            "unrelated_slice_activations": sum(1 for item in final.get("impacts", []) if item.get("affected_slice") in scenario.forbidden_slices),
            "model_calls": int(final_metrics.get("model_calls", 0)),
            "tool_calls": tool_calls,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "estimated_cost": round(estimated_cost, 6),
            "receipt_bytes_start": receipt_start,
            "receipt_bytes_end": receipt_end,
            "receipt_patch_bytes": int(final_metrics.get("temporal_patch_bytes", 0)),
            "receipt_patch_count": int(final_metrics.get("temporal_patch_count", 0)),
            "full_receipt_rebuilds": int(final_metrics.get("temporal_full_rebuilds", 0)),
            "freshness_model_calls": int(final_metrics.get("temporal_freshness_model_calls", 0)),
            "dependency_evaluations": int(final_metrics.get("temporal_dependencies_evaluated", 0)),
            "stale_facts": int(final_metrics.get("temporal_fields_marked_stale", 0)),
            "human_revalidations": int(final_metrics.get("temporal_one_tap_confirmations", 0) + final_metrics.get("temporal_bounded_update_flows", 0)),
            "machine_revalidations": int(final_metrics.get("temporal_fields_silently_refreshed", 0)),
            "state_versions": [state.get("version") for state in all_states],
            "latency": round(time.perf_counter() - started, 4),
            "interaction_latencies": latencies,
            "deterministic_latencies": deterministic_latencies,
            "model_backed_latencies": model_backed_latencies,
            "final_status": execution_state,
            "failure_codes": sorted(set(failures)),
            "events": events,
        }


def percentile(values: list[float], p: float) -> float | None:
    if not values:
        return None
    ordered = sorted(values)
    index = max(0, math.ceil(p * len(ordered)) - 1)
    return ordered[index]


def main() -> None:
    run_started = datetime.now(UTC)
    wall_started = time.perf_counter()
    results: list[dict[str, Any]] = []
    driver_errors: list[dict[str, str]] = []
    with ThreadPoolExecutor(max_workers=6) as pool:
        futures = {pool.submit(run_scenario, scenario): scenario for scenario in SCENARIOS}
        for future in as_completed(futures):
            scenario = futures[future]
            try:
                results.append(future.result())
            except Exception as exc:  # driver records provider/runtime failures as scenario data
                driver_errors.append({"scenario_id": scenario.id, "error": f"{type(exc).__name__}: {exc}"})
    results.sort(key=lambda item: item["scenario_id"])
    elapsed = time.perf_counter() - wall_started
    total_cost = sum(item["estimated_cost"] for item in results)
    scenario_latencies = [item["latency"] for item in results]
    deterministic_latencies = [value for item in results for value in item["deterministic_latencies"]]
    model_latencies = [value for item in results for value in item["model_backed_latencies"]]
    aggregate = {
        "frozen": {
            "base_url": BASE_URL,
            "revision": REVISION,
            "source_commit": SOURCE_COMMIT,
            "scenario_manifest_sha256": "computed-after-write",
            "start_timestamp": run_started.isoformat(),
            "starting_estimated_spend": STARTING_ESTIMATED_SPEND,
            "cost_assumption": {"input_usd_per_million": INPUT_USD_PER_MILLION, "output_usd_per_million": OUTPUT_USD_PER_MILLION, "tool_call_usd": TOOL_CALL_USD},
        },
        "budget": {"wall_seconds": round(elapsed, 3), "estimated_incremental_cost": round(total_cost, 6), "safe_cost_rail": SAFE_COST_RAIL, "hard_cost_rail": HARD_COST_RAIL},
        "counts": {"attempted": len(SCENARIOS), "completed": len(results), "driver_errors": len(driver_errors), "passed": sum(not item["failure_codes"] for item in results)},
        "economics": {
            "model_calls": sum(item["model_calls"] for item in results),
            "tool_calls": sum(item["tool_calls"] for item in results),
            "input_tokens": sum(item["input_tokens"] for item in results),
            "output_tokens": sum(item["output_tokens"] for item in results),
            "full_receipt_rebuilds": sum(item["full_receipt_rebuilds"] for item in results),
            "freshness_model_calls": sum(item["freshness_model_calls"] for item in results),
            "receipt_patch_bytes": sum(item["receipt_patch_bytes"] for item in results),
            "receipt_patch_count": sum(item["receipt_patch_count"] for item in results),
        },
        "freshness": {
            "dependency_evaluations": sum(item["dependency_evaluations"] for item in results),
            "stale_facts": sum(item["stale_facts"] for item in results),
            "human_revalidations": sum(item["human_revalidations"] for item in results),
            "machine_revalidations": sum(item["machine_revalidations"] for item in results),
        },
        "latency": {
            "median_scenario": statistics.median(scenario_latencies) if scenario_latencies else None,
            "p90_scenario": percentile(scenario_latencies, 0.90),
            "p95_scenario": percentile(scenario_latencies, 0.95),
            "maximum_scenario": max(scenario_latencies, default=None),
            "median_deterministic_interaction": statistics.median(deterministic_latencies) if deterministic_latencies else None,
            "p90_deterministic_interaction": percentile(deterministic_latencies, 0.90),
            "median_model_interaction": statistics.median(model_latencies) if model_latencies else None,
            "p90_model_interaction": percentile(model_latencies, 0.90),
            "maximum_model_interaction": max(model_latencies, default=None),
        },
        "driver_errors": driver_errors,
        "results": results,
    }
    output_dir = Path(__file__).with_name("output")
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / f"gauntlet-{run_started.strftime('%Y%m%dT%H%M%SZ')}.json"
    output.write_text(json.dumps(aggregate, indent=2, sort_keys=True), encoding="utf-8")
    print(output)
    print(json.dumps({key: aggregate[key] for key in ("budget", "counts", "economics", "freshness", "latency", "driver_errors")}, indent=2))


if __name__ == "__main__":
    main()
